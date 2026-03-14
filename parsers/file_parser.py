"""
File parser — extracts readable text from uploaded files.
Supports: .md, .txt, .html, .htm, .docx, .pdf
"""

from __future__ import annotations

import io
from pathlib import Path


def parse_file(file_bytes: bytes, filename: str) -> str:
    """Extract text content from an uploaded file."""
    ext = Path(filename).suffix.lower()

    parsers = {
        ".md": _parse_text,
        ".txt": _parse_text,
        ".html": _parse_html,
        ".htm": _parse_html,
        ".docx": _parse_docx,
        ".pdf": _parse_pdf,
    }

    parser = parsers.get(ext)
    if not parser:
        raise ValueError(
            f"Unsupported file type: {ext}. "
            f"Supported: {', '.join(parsers.keys())}"
        )

    return parser(file_bytes)


def _parse_text(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="replace")


def _parse_html(file_bytes: bytes) -> str:
    from bs4 import BeautifulSoup

    html = file_bytes.decode("utf-8", errors="replace")
    soup = BeautifulSoup(html, "html.parser")

    for tag in soup(["script", "style", "nav", "footer", "header", "noscript"]):
        tag.decompose()

    lines: list[str] = []

    for el in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "blockquote", "td", "th", "figcaption"]):
        text = el.get_text(strip=True)
        if not text:
            continue

        if el.name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(el.name[1])
            lines.append(f"{'#' * level} {text}")
        elif el.name == "li":
            lines.append(f"- {text}")
        elif el.name == "blockquote":
            lines.append(f"> {text}")
        else:
            lines.append(text)

    if not lines:
        return soup.get_text(separator="\n", strip=True)

    return "\n\n".join(lines)


def _parse_docx(file_bytes: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = (getattr(para.style, "name", None) or "").lower() if para.style else ""
        if "heading 1" in style_name:
            lines.append(f"# {text}")
        elif "heading 2" in style_name:
            lines.append(f"## {text}")
        elif "heading 3" in style_name:
            lines.append(f"### {text}")
        elif "heading 4" in style_name:
            lines.append(f"#### {text}")
        elif "list" in style_name:
            lines.append(f"- {text}")
        elif "quote" in style_name:
            lines.append(f"> {text}")
        else:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n\n".join(lines)


def _parse_pdf(file_bytes: bytes) -> str:
    import pdfplumber

    lines: list[str] = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                lines.append(text.strip())

    return "\n\n".join(lines)
